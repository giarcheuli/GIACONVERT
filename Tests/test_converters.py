#!/usr/bin/env python3
"""
Test existing GIACONVERT modules with various document types
This script tests our converters with different .docx files to ensure reliability.
"""

import os
import sys
from pathlib import Path
import tempfile
import importlib.util

def test_converter_modules():
    """Test all three converter modules with various document types"""
    
    print("🔍 Testing existing converter modules...")
    
    # Define our converter modules
    converters = {
        "basic": "giaconvert.py",
        "with_images": "giaconvert_with_images.py", 
        "complete": "giaconvert_complete.py"
    }
    
    base_dir = Path(__file__).parent
    
    # Check if all converter files exist
    for name, filename in converters.items():
        converter_path = base_dir / filename
        if not converter_path.exists():
            print(f"❌ {name} converter not found: {converter_path}")
            return False
        else:
            print(f"✅ {name} converter found: {filename}")
    
    # Check if test document exists
    test_doc = base_dir / "test_documents" / "sample_document.docx"
    if not test_doc.exists():
        print(f"❌ Test document not found: {test_doc}")
        print("📝 Creating a test document...")
        create_test_document()
    else:
        print(f"✅ Test document found: {test_doc}")
    
    # Test each converter
    test_results = {}
    
    for name, filename in converters.items():
        print(f"\n🧪 Testing {name} converter ({filename})...")
        
        try:
            # Import the converter module
            converter_path = base_dir / filename
            spec = importlib.util.spec_from_file_location(f"converter_{name}", converter_path)
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)
            
            # Check if WordToHTMLConverter class exists
            if hasattr(module, 'WordToHTMLConverter'):
                converter_class = module.WordToHTMLConverter
                print(f"✅ WordToHTMLConverter class found in {name}")
                
                # Test basic instantiation
                converter = converter_class()
                print(f"✅ {name} converter instantiated successfully")
                
                # Test with sample document if it exists
                if test_doc.exists():
                    with tempfile.TemporaryDirectory() as temp_dir:
                        output_path = Path(temp_dir) / f"test_output_{name}.html"
                        
                        try:
                            # Test conversion - our converters use convert_docx_to_html
                            html_output_file = str(output_path)
                            converter.convert_docx_to_html(str(test_doc), html_output_file)
                            
                            # Read the generated HTML content
                            if output_path.exists():
                                html_content = output_path.read_text(encoding='utf-8')
                            else:
                                html_content = ""
                            
                            # Basic validation
                            if html_content and len(html_content) > 100:
                                print(f"✅ {name} converter produced valid HTML output ({len(html_content)} chars)")
                                test_results[name] = {
                                    "status": "success",
                                    "output_length": len(html_content),
                                    "output_file": str(output_path)
                                }
                            else:
                                print(f"❌ {name} converter produced minimal output")
                                test_results[name] = {"status": "minimal_output"}
                                
                        except Exception as e:
                            print(f"❌ {name} converter failed during conversion: {e}")
                            test_results[name] = {"status": "conversion_error", "error": str(e)}
                
                else:
                    print(f"⚠️  {name} converter class OK, but no test document for conversion test")
                    test_results[name] = {"status": "no_test_doc"}
                    
            else:
                print(f"❌ WordToHTMLConverter class not found in {name}")
                test_results[name] = {"status": "no_class"}
                
        except Exception as e:
            print(f"❌ Failed to import {name} converter: {e}")
            test_results[name] = {"status": "import_error", "error": str(e)}
    
    # Print summary
    print("\n📋 Test Results Summary:")
    print("=" * 50)
    
    all_passed = True
    for name, result in test_results.items():
        status = result["status"]
        if status == "success":
            print(f"✅ {name.upper()}: PASSED - HTML output ({result['output_length']} chars)")
        elif status == "no_test_doc":
            print(f"⚠️  {name.upper()}: PARTIAL - Class OK, needs test document")
        else:
            print(f"❌ {name.upper()}: FAILED - {status}")
            if "error" in result:
                print(f"   Error: {result['error']}")
            all_passed = False
    
    print("=" * 50)
    
    if all_passed:
        print("🎉 All converter modules are working correctly!")
    else:
        print("⚠️  Some issues found. Need to fix before proceeding.")
    
    return test_results

def create_test_document():
    """Create a comprehensive test document for converter testing"""
    
    print("📝 Creating comprehensive test document...")
    
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('GIACONVERT Test Document', 0)
        
        # Add paragraphs with different styles
        doc.add_heading('1. Basic Text Testing', level=1)
        doc.add_paragraph('This is a simple paragraph to test basic text conversion.')
        doc.add_paragraph('This paragraph contains bold text and italic text.', style='Intense Quote')
        
        # Add a list
        doc.add_heading('2. List Testing', level=1)
        doc.add_paragraph('Bullet point 1', style='List Bullet')
        doc.add_paragraph('Bullet point 2', style='List Bullet')
        doc.add_paragraph('Bullet point 3', style='List Bullet')
        
        # Add a table
        doc.add_heading('3. Table Testing', level=1)
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Medium Shading 1 Accent 1'
        
        # Fill table with test data
        cells = table.rows[0].cells
        cells[0].text = 'Header 1'
        cells[1].text = 'Header 2'
        cells[2].text = 'Header 3'
        
        cells = table.rows[1].cells
        cells[0].text = 'Row 1, Col 1'
        cells[1].text = 'Row 1, Col 2'
        cells[2].text = 'Row 1, Col 3'
        
        cells = table.rows[2].cells
        cells[0].text = 'Row 2, Col 1'
        cells[1].text = 'Row 2, Col 2'
        cells[2].text = 'Row 2, Col 3'
        
        # Add more content for testing
        doc.add_heading('4. Formatting Testing', level=1)
        p = doc.add_paragraph()
        p.add_run('Normal text, ')
        p.add_run('bold text, ').bold = True
        p.add_run('italic text, ').italic = True
        p.add_run('underlined text.').underline = True
        
        # Create test_documents directory if it doesn't exist
        test_dir = Path(__file__).parent / "test_documents"
        test_dir.mkdir(exist_ok=True)
        
        # Save the document
        doc_path = test_dir / "sample_document.docx"
        doc.save(str(doc_path))
        
        print(f"✅ Test document created: {doc_path}")
        return str(doc_path)
        
    except ImportError:
        print("❌ python-docx not available for creating test document")
        return None
    except Exception as e:
        print(f"❌ Error creating test document: {e}")
        return None

if __name__ == "__main__":
    test_converter_modules()