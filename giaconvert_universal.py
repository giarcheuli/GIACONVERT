#!/usr/bin/env python3
"""
GIACONVERT Universal Converter
Supports both .doc and .docx files with comprehensive conversion features
"""

import os
import sys
import shutil
import zipfile
import base64
import re
from pathlib import Path
from typing import Optional, List, Dict, Any

# For .docx files
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# For .doc files
import docx2txt

# For HTML processing
from lxml import html, etree
try:
    from html import escape as html_escape
except ImportError:
    from cgi import escape as html_escape


class UniversalDocumentConverter:
    """Universal converter for both .doc and .docx files"""
    
    def __init__(self):
        self.image_counter = 0
        self.extracted_images = []
    
    def convert_doc_to_html(self, doc_path: str, html_path: str, extract_images: bool = False) -> Dict[str, Any]:
        """
        Convert .doc file to HTML using docx2txt
        
        Args:
            doc_path: Path to the .doc file
            html_path: Path where HTML file should be saved
            extract_images: Whether to extract images (limited support for .doc)
            
        Returns:
            Dictionary with conversion results
        """
        try:
            doc_path = Path(doc_path)
            html_path = Path(html_path)
            
            # Create output directory if it doesn't exist
            html_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Extract text from .doc file
            if extract_images:
                # Create temporary directory for images
                temp_dir = html_path.parent / "temp_images"
                temp_dir.mkdir(exist_ok=True)
                
                try:
                    # Extract text and images
                    text = docx2txt.process(str(doc_path), str(temp_dir))
                    
                    # Get extracted images
                    image_files = list(temp_dir.glob("*"))
                    images_dir = None
                    
                    if image_files:
                        # Create images directory
                        images_dir = html_path.parent / f"{html_path.stem}_images"
                        images_dir.mkdir(exist_ok=True)
                        
                        # Move images and track them
                        for i, img_file in enumerate(image_files):
                            if img_file.is_file():
                                new_name = f"image_{i+1}{img_file.suffix}"
                                new_path = images_dir / new_name
                                shutil.move(str(img_file), str(new_path))
                                self.extracted_images.append({
                                    'original_name': img_file.name,
                                    'new_name': new_name,
                                    'path': str(new_path)
                                })
                    
                    # Clean up temp directory
                    if temp_dir.exists():
                        shutil.rmtree(temp_dir)
                        
                except Exception as e:
                    print(f"Warning: Could not extract images from .doc file: {e}")
                    text = docx2txt.process(str(doc_path))
                    images_dir = None
            else:
                text = docx2txt.process(str(doc_path))
                images_dir = None
            
            # Convert text to HTML
            html_content = self._convert_text_to_html(text, doc_path.stem, images_dir)
            
            # Write HTML file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'html_path': str(html_path),
                'images_extracted': len(self.extracted_images),
                'images_dir': str(images_dir) if images_dir else None,
                'message': f'Successfully converted .doc file to HTML'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': f'Failed to convert .doc file: {str(e)}'
            }
    
    def convert_docx_to_html(self, docx_path: str, html_path: str, 
                           extract_images: bool = False, 
                           include_headers_footers: bool = False) -> Dict[str, Any]:
        """
        Convert .docx file to HTML with full feature support
        
        Args:
            docx_path: Path to the .docx file
            html_path: Path where HTML file should be saved
            extract_images: Whether to extract and embed images
            include_headers_footers: Whether to include headers and footers
            
        Returns:
            Dictionary with conversion results
        """
        try:
            docx_path = Path(docx_path)
            html_path = Path(html_path)
            
            # Create output directory if it doesn't exist
            html_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Load the document
            doc = Document(docx_path)
            
            # Reset counters
            self.image_counter = 0
            self.extracted_images = []

            # Prepare images directory for external mode (images are extracted inline during conversion)
            images_dir = None
            if extract_images:
                images_dir = html_path.parent / f"{html_path.stem}_images"
                images_dir.mkdir(exist_ok=True)
            
            # Convert document content
            html_content = self._convert_docx_content_to_html(
                doc, docx_path.stem, images_dir, include_headers_footers
            )
            
            # Write HTML file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'html_path': str(html_path),
                'images_extracted': len(self.extracted_images),
                'images_dir': str(images_dir) if images_dir and self.extracted_images else None,
                'message': f'Successfully converted .docx file to HTML'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': f'Failed to convert .docx file: {str(e)}'
            }
    
    def convert_document(self, input_path: str, output_path: str, 
                        mode: str = 'enhanced') -> Dict[str, Any]:
        """
        Universal converter method that handles both .doc and .docx files
        
        Args:
            input_path: Path to input document (.doc or .docx)
            output_path: Path for output HTML file
            mode: Conversion mode ('basic', 'enhanced', 'complete')
            
        Returns:
            Dictionary with conversion results
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        
        # Determine file type
        file_extension = input_path.suffix.lower()
        
        # Set conversion parameters based on mode
        extract_images = mode in ['enhanced', 'complete']
        include_headers_footers = mode == 'complete'
        
        if file_extension == '.docx':
            return self.convert_docx_to_html(
                str(input_path), 
                str(output_path), 
                extract_images=extract_images,
                include_headers_footers=include_headers_footers
            )
        elif file_extension == '.doc':
            return self.convert_doc_to_html(
                str(input_path), 
                str(output_path), 
                extract_images=extract_images
            )
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_extension}',
                'message': f'Only .doc and .docx files are supported'
            }
    
    def _convert_text_to_html(self, text: str, title: str, images_dir: Optional[Path]) -> str:
        """Convert plain text to HTML with basic formatting"""
        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
            color: #333;
        }}
        .header {{
            border-bottom: 2px solid #333;
            margin-bottom: 20px;
            padding-bottom: 10px;
        }}
        .content {{
            max-width: 800px;
        }}
        p {{
            margin-bottom: 15px;
        }}
        .image {{
            max-width: 100%;
            height: auto;
            margin: 20px 0;
            border: 1px solid #ddd;
            padding: 5px;
        }}
        .note {{
            background-color: #f0f8ff;
            padding: 10px;
            border-left: 4px solid #007acc;
            margin: 20px 0;
            font-style: italic;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
        <p><em>Converted from .doc format</em></p>
    </div>
    <div class="content">
"""
        
        # Add paragraphs
        for paragraph in paragraphs:
            if paragraph.strip():
                # Simple formatting detection
                lines = paragraph.split('\n')
                for line in lines:
                    if line.strip():
                        # Check if it looks like a heading
                        if len(line) < 100 and (line.isupper() or line.startswith('Chapter') or line.startswith('Section')):
                            html_content += f"        <h2>{html_escape(line.strip())}</h2>\n"
                        else:
                            html_content += f"        <p>{html_escape(line.strip())}</p>\n"
        
        # Add images if any were extracted
        if images_dir and self.extracted_images:
            html_content += f"""
        <div class="note">
            <strong>Note:</strong> This document contained {len(self.extracted_images)} image(s) 
            which have been extracted to the <code>{images_dir.name}/</code> folder.
        </div>
"""
            
            for img in self.extracted_images:
                rel_path = f"{images_dir.name}/{img['new_name']}"
                html_content += f'        <img src="{rel_path}" alt="{img["original_name"]}" class="image" />\n'
        
        html_content += """    </div>
</body>
</html>"""
        
        return html_content
    
    def _get_image_extension(self, image_data: bytes) -> str:
        """Determine image file extension from binary magic bytes"""
        if image_data.startswith(b'\x89PNG'):
            return 'png'
        elif image_data.startswith(b'\xff\xd8\xff'):
            return 'jpg'
        elif image_data.startswith(b'GIF'):
            return 'gif'
        elif image_data.startswith(b'BM'):
            return 'bmp'
        return 'png'

    def _get_run_style(self, run) -> str:
        """Extract inline CSS styling from a docx run"""
        styles = []

        if run.bold:
            styles.append('font-weight:bold')
        if run.italic:
            styles.append('font-style:italic')
        if run.underline:
            styles.append('text-decoration:underline')

        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                if hasattr(rgb, 'red'):
                    r, g, b = rgb.red, rgb.green, rgb.blue
                else:
                    r, g, b = rgb[0], rgb[1], rgb[2]
                styles.append(f'color:rgb({r},{g},{b})')
        except (AttributeError, IndexError, TypeError):
            pass

        try:
            if run.font.size:
                size_px = int(run.font.size.pt * 1.33)
                styles.append(f'font-size:{size_px}px')
        except (AttributeError, TypeError):
            pass

        if run.font.name:
            styles.append(f"font-family:'{run.font.name}',sans-serif")

        return ';'.join(styles)

    def _extract_paragraph_images(self, doc, paragraph, images_dir: Optional[Path]) -> str:
        """
        Walk paragraph runs and extract any inline images (w:drawing elements).
        Returns HTML <img> tags for all found images, placed at their document position.
        """
        html_parts = []
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }

        for run in paragraph.runs:
            for drawing in run._element.xpath('.//w:drawing'):
                try:
                    blips = drawing.xpath('.//a:blip[@r:embed]', namespaces=namespaces)
                    for blip in blips:
                        rel_id = blip.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                        )
                        if not rel_id:
                            continue
                        try:
                            image_data = doc.part.rels[rel_id].target_part.blob
                        except (KeyError, AttributeError):
                            continue

                        self.image_counter += 1
                        ext = self._get_image_extension(image_data)

                        if images_dir is not None:
                            # Save as external file
                            filename = f'image_{self.image_counter:03d}.{ext}'
                            img_path = images_dir / filename
                            with open(img_path, 'wb') as f:
                                f.write(image_data)
                            rel_path = f'{images_dir.name}/{filename}'
                            html_parts.append(
                                f'<img src="{rel_path}" alt="Image {self.image_counter}" '
                                f'style="max-width:100%;height:auto;" />'
                            )
                            self.extracted_images.append({
                                'original_name': filename,
                                'new_name': filename,
                                'path': str(img_path),
                            })
                        else:
                            # Embed as base64
                            mime = f"image/{'jpeg' if ext == 'jpg' else ext}"
                            b64 = base64.b64encode(image_data).decode('utf-8')
                            src = f'data:{mime};base64,{b64}'
                            html_parts.append(
                                f'<img src="{src}" alt="Image {self.image_counter}" '
                                f'style="max-width:100%;height:auto;" />'
                            )
                            self.extracted_images.append({
                                'original_name': f'image_{self.image_counter}.{ext}',
                                'new_name': f'image_{self.image_counter}.{ext}',
                                'path': None,
                            })
                except Exception:
                    pass

        return ''.join(html_parts)

    def _convert_paragraph(self, doc, paragraph, images_dir: Optional[Path]) -> str:
        """Convert a single docx paragraph to an HTML element, including inline images."""
        # Collect run text with inline styling
        text_parts = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            text = html_escape(text)
            style = self._get_run_style(run)
            if style:
                text_parts.append(f'<span style="{style}">{text}</span>')
            else:
                text_parts.append(text)

        image_html = self._extract_paragraph_images(doc, paragraph, images_dir)

        content = ''.join(text_parts)
        if image_html:
            content = (content + image_html) if content else image_html

        if not content:
            return '<br/>'

        # Heading styles
        style_name = paragraph.style.name
        if style_name.startswith('Heading'):
            level_str = style_name.replace('Heading ', '')
            try:
                level = min(int(level_str), 6)
            except ValueError:
                level = 2
            return f'<h{level}>{content}</h{level}>'

        # Paragraph alignment
        _align_map = {
            WD_PARAGRAPH_ALIGNMENT.CENTER: 'center',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: 'right',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'justify',
        }
        alignment = _align_map.get(paragraph.alignment, '')
        if alignment:
            return f'<p style="text-align:{alignment}">{content}</p>'
        return f'<p>{content}</p>'

    def _convert_table_to_html(self, doc, table, images_dir: Optional[Path]) -> str:
        """Convert a docx table to HTML, with rich cell content."""
        rows_html = []
        for i, row in enumerate(table.rows):
            cells_html = []
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                cell_parts = []
                for para in cell.paragraphs:
                    cell_parts.append(self._convert_paragraph(doc, para, images_dir))
                cell_content = ''.join(cell_parts) if cell_parts else '&nbsp;'
                cells_html.append(f'<{tag}>{cell_content}</{tag}>')
            rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')
        return '<table>\n' + '\n'.join(rows_html) + '\n</table>'

    def _extract_headers_footers_html(self, doc, part: str, images_dir: Optional[Path]) -> str:
        """
        Extract real header or footer content from all document sections.
        `part` is either 'header' or 'footer'.
        Returns combined HTML string, or empty string if nothing found.
        """
        parts_html = []
        try:
            for section in doc.sections:
                section_part = section.header if part == 'header' else section.footer
                if section_part is None:
                    continue
                for para in section_part.paragraphs:
                    if para.text.strip():
                        parts_html.append(self._convert_paragraph(doc, para, images_dir))
        except Exception as e:
            print(f"Warning: Could not extract {part}: {e}")
        return '\n'.join(parts_html)

    def _convert_docx_content_to_html(self, doc, title: str, images_dir: Optional[Path],
                                      include_headers_footers: bool) -> str:
        """Convert .docx document content to HTML with inline images and real headers/footers."""

        header_footer_css = ''
        if include_headers_footers:
            header_footer_css = '''
        .document-header {
            background: #f8f9fa;
            border-bottom: 2px solid #e9ecef;
            padding: 15px 40px;
            margin-bottom: 20px;
        }
        .document-footer {
            background: #f8f9fa;
            border-top: 2px solid #e9ecef;
            padding: 15px 40px;
            margin-top: 20px;
        }
        @media print {
            .document-header { position: running(header); background: white !important; border: none !important; }
            .document-footer { position: running(footer); background: white !important; border: none !important; }
        }'''

        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="en">',
            '<head>',
            '    <meta charset="UTF-8">',
            '    <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            f'    <title>{html_escape(title)}</title>',
            '    <style>',
            '        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; color: #333; }',
            '        p { margin-bottom: 15px; }',
            '        img { max-width: 100%; height: auto; margin: 10px 0; display: block; }',
            '        table { border-collapse: collapse; width: 100%; margin: 20px 0; }',
            '        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }',
            '        th { background-color: #f2f2f2; }',
            header_footer_css,
            '    </style>',
            '</head>',
            '<body>',
        ]

        # Real header content
        if include_headers_footers:
            headers_html = self._extract_headers_footers_html(doc, 'header', images_dir)
            if headers_html:
                html_parts.append('<header class="document-header">')
                html_parts.append(headers_html)
                html_parts.append('</header>')

        html_parts.append('<main class="document-content">')

        # Iterate over body elements in document order to preserve layout
        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'p':
                for para in doc.paragraphs:
                    if para._element is element:
                        html_parts.append(self._convert_paragraph(doc, para, images_dir))
                        break
            elif tag == 'tbl':
                for table in doc.tables:
                    if table._element is element:
                        html_parts.append(self._convert_table_to_html(doc, table, images_dir))
                        break

        html_parts.append('</main>')

        # Real footer content
        if include_headers_footers:
            footers_html = self._extract_headers_footers_html(doc, 'footer', images_dir)
            if footers_html:
                html_parts.append('<footer class="document-footer">')
                html_parts.append(footers_html)
                html_parts.append('</footer>')

        html_parts.extend(['</body>', '</html>'])
        return '\n'.join(html_parts)


def main():
    """Command line interface for testing"""
    if len(sys.argv) < 3:
        print("Usage: python giaconvert_universal.py <input_file> <output_file> [mode]")
        print("Modes: basic, enhanced, complete")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    mode = sys.argv[3] if len(sys.argv) > 3 else 'enhanced'
    
    converter = UniversalDocumentConverter()
    result = converter.convert_document(input_file, output_file, mode)
    
    if result['success']:
        print(f"✅ {result['message']}")
        if result.get('images_extracted', 0) > 0:
            print(f"📷 Extracted {result['images_extracted']} images")
    else:
        print(f"❌ {result['message']}")
        sys.exit(1)


if __name__ == "__main__":
    main()