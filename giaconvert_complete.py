#!/usr/bin/env python3
"""
GIACONVERT with Headers, Footers, and Image Support - Complete Version
Converts .docx files to HTML format while preserving formatting, structure, images, headers, and footers.
"""

import os
import sys
import click
import base64
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.document import Document as DocumentType
import xml.etree.ElementTree as ET
from PIL import Image
import io


class WordToHTMLConverter:
    def __init__(self, image_mode='external', optimize_images=False, headers_footers='include'):
        self.converted_count = 0
        self.error_count = 0
        self.errors = []
        self.image_mode = image_mode  # 'external', 'inline', or 'skip'
        self.optimize_images = optimize_images
        self.headers_footers = headers_footers  # 'include', 'skip', or 'print-only'
        self.image_counter = 0

    def convert_paragraph_alignment(self, alignment):
        """Convert docx alignment to CSS text-align"""
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: 'left',
            WD_PARAGRAPH_ALIGNMENT.CENTER: 'center',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: 'right',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'justify',
        }
        return alignment_map.get(alignment, 'left')

    def get_run_style(self, run):
        """Extract styling from a run and return CSS style string"""
        styles = []
        
        if run.bold:
            styles.append("font-weight: bold")
        if run.italic:
            styles.append("font-style: italic")
        if run.underline:
            styles.append("text-decoration: underline")
        
        # Font color
        if run.font.color and run.font.color.rgb:
            try:
                rgb = run.font.color.rgb
                if hasattr(rgb, 'red'):
                    r, g, b = rgb.red, rgb.green, rgb.blue
                else:
                    r, g, b = rgb[0], rgb[1], rgb[2]
                styles.append(f"color: rgb({r}, {g}, {b})")
            except (AttributeError, IndexError, TypeError):
                pass
        
        # Font size
        if run.font.size:
            try:
                size_px = int(run.font.size.pt * 1.33)
                styles.append(f"font-size: {size_px}px")
            except (AttributeError, TypeError):
                pass
        
        # Font name
        if run.font.name:
            styles.append(f"font-family: '{run.font.name}', sans-serif")
        
        return "; ".join(styles) if styles else ""

    def create_images_directory(self, html_path):
        """Create directory for external images"""
        images_dir = html_path.parent / f"{html_path.stem}_images"
        images_dir.mkdir(exist_ok=True)
        return images_dir

    def optimize_image(self, image_data, max_width=1200, max_height=800, quality=85):
        """Optimize image size and quality"""
        try:
            img = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if necessary (for JPEG output)
            if img.mode in ('RGBA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            
            # Resize if too large
            if img.width > max_width or img.height > max_height:
                img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            # Save optimized image
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=quality, optimize=True)
            return output.getvalue()
            
        except Exception as e:
            # If optimization fails, return original
            return image_data

    def extract_image_data(self, doc, image_rel_id):
        """Extract image data from document relationships"""
        try:
            # Get the image part from document relationships
            image_part = doc.part.rels[image_rel_id].target_part
            return image_part.blob
        except (KeyError, AttributeError):
            return None

    def get_image_extension(self, image_data):
        """Determine image file extension from binary data"""
        if image_data.startswith(b'\x89PNG'):
            return 'png'
        elif image_data.startswith(b'\xff\xd8\xff'):
            return 'jpg'
        elif image_data.startswith(b'GIF'):
            return 'gif'
        elif image_data.startswith(b'BM'):
            return 'bmp'
        else:
            return 'png'  # Default fallback

    def convert_image_to_base64(self, image_data):
        """Convert image data to base64 string"""
        extension = self.get_image_extension(image_data)
        mime_type = f"image/{extension.replace('jpg', 'jpeg')}"
        b64_data = base64.b64encode(image_data).decode('utf-8')
        return f"data:{mime_type};base64,{b64_data}"

    def process_paragraph_images(self, paragraph, html_path, images_dir):
        """Process images within a paragraph and return HTML"""
        if self.image_mode == 'skip':
            return ""
        
        html_parts = []
        
        # Check for inline shapes (images)
        for run in paragraph.runs:
            # Look for drawing elements in the run
            for drawing in run._element.xpath('.//w:drawing'):
                try:
                    # Find image relationships
                    blips = drawing.xpath('.//a:blip[@r:embed]', 
                                        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                                                  'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
                    
                    for blip in blips:
                        rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rel_id:
                            image_data = self.extract_image_data(paragraph._parent, rel_id)
                            if image_data:
                                self.image_counter += 1
                                
                                if self.image_mode == 'external':
                                    # Save as external file
                                    if self.optimize_images:
                                        image_data = self.optimize_image(image_data)
                                    
                                    extension = self.get_image_extension(image_data)
                                    image_filename = f"image_{self.image_counter:03d}.{extension}"
                                    image_path = images_dir / image_filename
                                    
                                    with open(image_path, 'wb') as f:
                                        f.write(image_data)
                                    
                                    # Relative path from HTML to image
                                    relative_path = f"{images_dir.name}/{image_filename}"
                                    html_parts.append(f'<img src="{relative_path}" alt="Image {self.image_counter}" style="max-width: 100%; height: auto;"/>')
                                
                                elif self.image_mode == 'inline':
                                    # Embed as base64
                                    if self.optimize_images:
                                        image_data = self.optimize_image(image_data)
                                    
                                    base64_src = self.convert_image_to_base64(image_data)
                                    html_parts.append(f'<img src="{base64_src}" alt="Image {self.image_counter}" style="max-width: 100%; height: auto;"/>')
                
                except Exception as e:
                    self.errors.append(f"Error processing image: {str(e)}")
        
        return ''.join(html_parts)

    def convert_paragraph_to_html(self, paragraph, html_path=None, images_dir=None):
        """Convert a docx paragraph to HTML with image support"""
        # First, check for images
        image_html = ""
        if html_path and images_dir:
            image_html = self.process_paragraph_images(paragraph, html_path, images_dir)
        
        # If paragraph is empty but has images, return just the images
        if not paragraph.text.strip() and image_html:
            return f'<p>{image_html}</p>'
        
        # If paragraph is completely empty, return line break
        if not paragraph.text.strip() and not image_html:
            return "<br/>"
        
        # Get paragraph alignment
        alignment = self.convert_paragraph_alignment(paragraph.alignment)
        
        # Start building the HTML
        html_content = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Escape HTML characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Get run styling
            style = self.get_run_style(run)
            
            if style:
                html_content.append(f'<span style="{style}">{text}</span>')
            else:
                html_content.append(text)
        
        # Combine text and images
        content = ''.join(html_content)
        if image_html:
            content = content + image_html if content else image_html
        
        # Wrap in paragraph tag with alignment
        if alignment != 'left':
            return f'<p style="text-align: {alignment}">{content}</p>'
        else:
            return f'<p>{content}</p>'

    def convert_table_to_html(self, table):
        """Convert a docx table to HTML"""
        html = ['<table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse;">']
        
        for row in table.rows:
            html.append('<tr>')
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_content.append(self.convert_paragraph_to_html(paragraph))
                
                cell_html = ''.join(cell_content) if cell_content else '&nbsp;'
                html.append(f'<td>{cell_html}</td>')
            html.append('</tr>')
        
        html.append('</table>')
        return ''.join(html)

    def extract_headers_footers(self, doc):
        """Extract headers and footers from document sections"""
        headers_footers = {
            'headers': [],
            'footers': []
        }
        
        if self.headers_footers == 'skip':
            return headers_footers
        
        try:
            for section in doc.sections:
                # Extract header
                if section.header:
                    header_content = []
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():  # Only non-empty paragraphs
                            header_content.append(self.convert_paragraph_to_html(paragraph))
                    
                    if header_content:
                        headers_footers['headers'].extend(header_content)
                
                # Extract footer
                if section.footer:
                    footer_content = []
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():  # Only non-empty paragraphs
                            footer_content.append(self.convert_paragraph_to_html(paragraph))
                    
                    if footer_content:
                        headers_footers['footers'].extend(footer_content)
                        
        except Exception as e:
            self.errors.append(f"Error extracting headers/footers: {str(e)}")
        
        return headers_footers

    def generate_header_footer_css(self):
        """Generate CSS for headers and footers"""
        if self.headers_footers == 'skip':
            return ""
            
        css = []
        
        if self.headers_footers == 'include':
            css.extend([
                '.document-header {',
                '  background: #f8f9fa;',
                '  border-bottom: 2px solid #e9ecef;',
                '  padding: 15px 40px;',
                '  margin-bottom: 20px;',
                '}',
                '.document-footer {',
                '  background: #f8f9fa;',
                '  border-top: 2px solid #e9ecef;',
                '  padding: 15px 40px;',
                '  margin-top: 20px;',
                '}',
            ])
        
        if self.headers_footers in ['include', 'print-only']:
            css.extend([
                '@media print {',
                '  @page {',
                '    margin-top: 2cm;',
                '    margin-bottom: 2cm;',
                '  }',
                '  .document-header {',
                '    position: running(header);',
                '    background: white !important;',
                '    border: none !important;',
                '    margin: 0 !important;',
                '    padding: 10px 0 !important;',
                '  }',
                '  .document-footer {',
                '    position: running(footer);',
                '    background: white !important;',
                '    border: none !important;',
                '    margin: 0 !important;',
                '    padding: 10px 0 !important;',
                '  }',
                '  .no-print { display: none !important; }',
                '}',
            ])
            
        if self.headers_footers == 'print-only':
            css.extend([
                '@media screen {',
                '  .document-header, .document-footer { display: none; }',
                '}',
            ])
        
        return '\n'.join(css)

    def convert_docx_to_html(self, docx_path, html_path):
        """Convert a single docx file to HTML with image and header/footer support"""
        try:
            doc = Document(docx_path)
            self.image_counter = 0  # Reset counter for each document
            
            # Create images directory if using external mode
            images_dir = None
            if self.image_mode == 'external':
                images_dir = self.create_images_directory(html_path)
            
            # Extract headers and footers
            headers_footers = self.extract_headers_footers(doc)
            
            # Start building HTML
            html_parts = [
                '<!DOCTYPE html>',
                '<html>',
                '<head>',
                '<meta charset="UTF-8">',
                '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
                f'<title>{Path(docx_path).stem}</title>',
                '<style>',
                'body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }',
                'table { margin: 20px 0; width: 100%; }',
                'p { margin: 10px 0; }',
                'img { margin: 10px 0; display: block; }',
                self.generate_header_footer_css(),
                '</style>',
                '</head>',
                '<body>',
            ]
            
            # Add header if present and not skipped
            if headers_footers['headers'] and self.headers_footers != 'skip':
                html_parts.append('<header class="document-header">')
                html_parts.extend(headers_footers['headers'])
                html_parts.append('</header>')
            
            # Add main content wrapper
            html_parts.append('<main class="document-content">')
            
            # Convert document content
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    for para in doc.paragraphs:
                        if para._element == element:
                            html_parts.append(self.convert_paragraph_to_html(para, html_path, images_dir))
                            break
                elif element.tag.endswith('tbl'):  # Table
                    for table in doc.tables:
                        if table._element == element:
                            html_parts.append(self.convert_table_to_html(table))
                            break
            
            # Close main content wrapper
            html_parts.append('</main>')
            
            # Add footer if present and not skipped
            if headers_footers['footers'] and self.headers_footers != 'skip':
                html_parts.append('<footer class="document-footer">')
                html_parts.extend(headers_footers['footers'])
                html_parts.append('</footer>')
            
            html_parts.extend(['</body>', '</html>'])
            
            # Write HTML file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_parts))
            
            return True
            
        except Exception as e:
            self.errors.append(f"Error converting {docx_path}: {str(e)}")
            return False

    def find_word_documents(self, directory):
        """Find all .docx files in directory and subdirectories"""
        word_files = []
        directory = Path(directory)
        
        for file_path in directory.rglob('*.docx'):
            if not file_path.name.startswith('~$'):
                word_files.append(file_path)
        
        return word_files

    def convert_directory(self, directory):
        """Convert all Word documents in directory and subdirectories"""
        directory = Path(directory)
        
        if not directory.exists():
            click.echo(f"Error: Directory '{directory}' does not exist.", err=True)
            return False
        
        if not directory.is_dir():
            click.echo(f"Error: '{directory}' is not a directory.", err=True)
            return False
        
        word_files = self.find_word_documents(directory)
        
        if not word_files:
            click.echo(f"No Word documents (.docx) found in '{directory}' and its subdirectories.")
            return True
        
        click.echo(f"Found {len(word_files)} Word document(s) to convert...")
        if self.image_mode != 'skip':
            click.echo(f"Image handling: {self.image_mode}")
        click.echo(f"Headers/Footers: {self.headers_footers}")
        
        for docx_path in word_files:
            html_path = docx_path.with_suffix('.html')
            
            click.echo(f"Converting: {docx_path.relative_to(directory)}")
            
            if self.convert_docx_to_html(docx_path, html_path):
                self.converted_count += 1
                click.echo(f"  ✓ Converted to: {html_path.relative_to(directory)}")
                if self.image_counter > 0 and self.image_mode != 'skip':
                    click.echo(f"  📷 Images processed: {self.image_counter}")
            else:
                self.error_count += 1
                click.echo(f"  ✗ Failed to convert", err=True)
        
        return True


@click.command()
@click.argument('directory', type=click.Path(exists=True, file_okay=False, dir_okay=True, readable=True))
@click.option('--verbose', '-v', is_flag=True, help='Show detailed output')
@click.option('--images', type=click.Choice(['external', 'inline', 'skip']), default='external',
              help='How to handle images: external (separate files), inline (base64), skip (ignore)')
@click.option('--optimize-images', is_flag=True, help='Optimize images for web (resize and compress)')
@click.option('--headers-footers', type=click.Choice(['include', 'skip', 'print-only']), default='include',
              help='How to handle headers and footers: include (show on screen and print), skip (ignore), print-only (only for print)')
def main(directory, verbose, images, optimize_images, headers_footers):
    """
    Convert Word documents (.docx) to HTML format with full support for images, headers, and footers.
    
    DIRECTORY: Path to the directory containing Word documents to convert.
    The tool will search recursively through all subdirectories.
    """
    click.echo("🔄 GIACONVERT - Complete Word to HTML Converter")
    click.echo("=" * 55)
    
    converter = WordToHTMLConverter(
        image_mode=images, 
        optimize_images=optimize_images,
        headers_footers=headers_footers
    )
    
    success = converter.convert_directory(directory)
    
    if not success:
        sys.exit(1)
    
    # Show summary
    click.echo("\n" + "=" * 55)
    click.echo("📊 Conversion Summary:")
    click.echo(f"  ✅ Successfully converted: {converter.converted_count}")
    click.echo(f"  ❌ Failed conversions: {converter.error_count}")
    
    if converter.errors and (verbose or converter.error_count > 0):
        click.echo("\n🚨 Errors encountered:")
        for error in converter.errors:
            click.echo(f"  • {error}")
    
    if converter.converted_count > 0:
        click.echo(f"\n🎉 Conversion completed!")
        if images == 'external':
            click.echo("📁 Images saved as separate files in *_images folders")
        elif images == 'inline':
            click.echo("📎 Images embedded directly in HTML files")
        elif images == 'skip':
            click.echo("🚫 Images were skipped")
        
        if headers_footers == 'include':
            click.echo("📄 Headers and footers included in HTML output")
        elif headers_footers == 'print-only':
            click.echo("🖨️ Headers and footers included for print only")
        elif headers_footers == 'skip':
            click.echo("🚫 Headers and footers were skipped")


if __name__ == '__main__':
    main()