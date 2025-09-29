#!/usr/bin/env python3
"""
Test script to create a sample Word document with images for testing the enhanced converter
"""

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from PIL import Image, ImageDraw

def create_test_images():
    """Create simple test images"""
    images_created = []
    
    # Create a simple colored rectangle image
    img1 = Image.new('RGB', (300, 200), color='lightblue')
    draw1 = ImageDraw.Draw(img1)
    draw1.text((10, 10), "Test Image 1", fill='black')
    img1_path = 'test_image_1.png'
    img1.save(img1_path)
    images_created.append(img1_path)
    
    # Create another test image
    img2 = Image.new('RGB', (250, 150), color='lightgreen')
    draw2 = ImageDraw.Draw(img2)
    draw2.text((10, 10), "Test Image 2", fill='black')
    img2_path = 'test_image_2.png'
    img2.save(img2_path)
    images_created.append(img2_path)
    
    return images_created

def create_test_document_with_images():
    """Create a test Word document with various formatting and images"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Test Document with Images', 0)
    
    # Introduction paragraph
    intro = doc.add_paragraph('This is a ')
    intro.add_run('test document').bold = True
    intro.add_run(' created to verify the ')
    italic_run = intro.add_run('Word to HTML converter with image support')
    italic_run.italic = True
    italic_run.font.color.rgb = RGBColor(0, 100, 200)
    intro.add_run(' functionality.')
    
    # Create test images
    print("Creating test images...")
    test_images = create_test_images()
    
    # Add first image
    if os.path.exists(test_images[0]):
        doc.add_paragraph('Here is the first test image:')
        doc.add_picture(test_images[0], width=Inches(3))
    
    # Formatted paragraph with different alignments
    center_para = doc.add_paragraph('This paragraph is centered')
    center_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add second image
    if os.path.exists(test_images[1]):
        doc.add_paragraph('And here is the second test image:')
        doc.add_picture(test_images[1], width=Inches(2.5))
    
    # List-like content
    doc.add_paragraph('Features being tested:')
    doc.add_paragraph('• Bold and italic text formatting', style='List Bullet')
    doc.add_paragraph('• Color changes in text', style='List Bullet')
    doc.add_paragraph('• Different paragraph alignments', style='List Bullet')
    doc.add_paragraph('• Table creation and formatting', style='List Bullet')
    doc.add_paragraph('• Image insertion and conversion', style='List Bullet')
    
    # Create a simple table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Fill table with data
    headers = ['Feature', 'Status', 'Notes']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # Sample data
    data = [
        ['Text Formatting', '✅ Working', 'Bold, italic, colors'],
        ['Image Support', '🔄 Testing', 'PNG, JPEG formats']
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.cell(row_idx, col_idx).text = cell_data
    
    # Final paragraph
    conclusion = doc.add_paragraph()
    conclusion.add_run('This document demonstrates various ').font.size = 14
    underlined_run = conclusion.add_run('formatting and image features')
    underlined_run.underline = True
    underlined_run.font.color.rgb = RGBColor(255, 0, 0)
    conclusion.add_run(' that should be preserved in the HTML conversion.')
    
    # Clean up test images
    for img_path in test_images:
        try:
            os.remove(img_path)
        except OSError:
            pass
    
    return doc

def main():
    # Create test directory
    test_dir = 'test_documents'
    os.makedirs(test_dir, exist_ok=True)
    
    # Create and save test document with images
    doc = create_test_document_with_images()
    doc_path = os.path.join(test_dir, 'sample_document_with_images.docx')
    doc.save(doc_path)
    
    print(f"✅ Test document with images created: {doc_path}")
    print("You can now test GIACONVERT with image support using:")
    print(f"python3 giaconvert_with_images.py {test_dir} --images external")
    print(f"python3 giaconvert_with_images.py {test_dir} --images inline")
    print(f"python3 giaconvert_with_images.py {test_dir} --images skip")

if __name__ == '__main__':
    main()