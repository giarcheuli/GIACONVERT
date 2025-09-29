#!/usr/bin/env python3
"""
Create a test .doc file for testing universal converter
Note: This creates a .docx file but renames it to .doc for basic testing
In practice, users would have real .doc files from older Word versions
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches

def create_test_doc_file():
    """Create a test .doc file (actually .docx renamed to .doc)"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Test Document (.doc format)', 0)
    
    # Add some content
    doc.add_paragraph('This is a test document to verify that our GIACONVERT application can handle both .doc and .docx files.')
    
    # Add a heading
    doc.add_heading('Features Tested', level=1)
    
    # Add bullet points
    p = doc.add_paragraph('Text formatting and paragraphs', style='List Bullet')
    doc.add_paragraph('Headings and structure', style='List Bullet')
    doc.add_paragraph('Basic document content', style='List Bullet')
    
    # Add another section
    doc.add_heading('Legacy Format Support', level=1)
    doc.add_paragraph('This document demonstrates that GIACONVERT now supports both:')
    doc.add_paragraph('• Modern .docx files (Office 2007+)', style='List Bullet')
    doc.add_paragraph('• Legacy .doc files (Office 97-2003)', style='List Bullet')
    
    # Add a table
    doc.add_heading('Test Table', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File Type'
    hdr_cells[1].text = 'Extension'
    hdr_cells[2].text = 'Support Status'
    
    # Add data rows
    row_cells = table.add_row().cells
    row_cells[0].text = 'Word Document'
    row_cells[1].text = '.docx'
    row_cells[2].text = '✅ Supported'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Legacy Word Document'
    row_cells[1].text = '.doc'
    row_cells[2].text = '✅ Supported'
    
    # Add final paragraph
    doc.add_paragraph('\nThis test file helps verify that the universal converter can process legacy Word documents alongside modern ones.')
    
    # Save the document
    test_dir = Path('test_documents')
    test_dir.mkdir(exist_ok=True)
    
    # Save as .docx first
    docx_path = test_dir / 'legacy_test_document.docx'
    doc.save(str(docx_path))
    
    # Create a copy as .doc (for testing purposes)
    doc_path = test_dir / 'legacy_test_document.doc'
    
    # Copy the content (this is a workaround since we can't create real .doc files with python-docx)
    import shutil
    shutil.copy2(str(docx_path), str(doc_path))
    
    print(f"✅ Test documents created:")
    print(f"   - {docx_path}")
    print(f"   - {doc_path}")
    print(f"\nNote: The .doc file is actually a .docx file renamed for testing.")
    print(f"In practice, users will have real .doc files from older Word versions.")
    
    return docx_path, doc_path

if __name__ == "__main__":
    create_test_doc_file()