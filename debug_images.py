#!/usr/bin/env python3
"""
Debug script to check image detection in Word documents
"""

from docx import Document
import sys

def debug_document_images(docx_path):
    """Debug image content in a Word document"""
    print(f"Debugging images in: {docx_path}")
    
    doc = Document(docx_path)
    
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    image_count = 0
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: '{paragraph.text[:50]}...'")
        
        # Check for runs with drawings
        for j, run in enumerate(paragraph.runs):
            # Look for drawing elements in the run
            drawings = run._element.xpath('.//w:drawing')
            if drawings:
                print(f"  Run {j} has {len(drawings)} drawing(s)")
                image_count += len(drawings)
                
                for k, drawing in enumerate(drawings):
                    # Find image relationships
                    blips = drawing.xpath('.//a:blip[@r:embed]', 
                                        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                                                  'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
                    print(f"    Drawing {k} has {len(blips)} blip(s)")
                    
                    for blip in blips:
                        rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        print(f"      Relationship ID: {rel_id}")
    
    print(f"\nTotal images found: {image_count}")
    
    # Also check document relationships
    print("\nDocument relationships:")
    try:
        for rel_id, rel in doc.part.rels.items():
            if 'image' in rel.reltype:
                print(f"  {rel_id}: {rel.reltype} -> {rel.target_part}")
    except Exception as e:
        print(f"Error checking relationships: {e}")

if __name__ == '__main__':
    if len(sys.argv) > 1:
        debug_document_images(sys.argv[1])
    else:
        debug_document_images('test_documents/sample_document_with_images.docx')