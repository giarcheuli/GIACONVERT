#!/usr/bin/env python3
"""
GIACONVERT - Word to HTML Converter CLI App
Converts .docx files to HTML format while preserving formatting and structure.
"""

import os
import sys
import click
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET


class WordToHTMLConverter:
    def __init__(self):
        self.converted_count = 0
        self.error_count = 0
        self.errors = []

    def convert_paragraph_alignment(self, alignment):
        """Convert docx alignment to CSS text-align"""
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: 'left',
            WD_PARAGRAPH_ALIGNMENT.CENTER: 'center',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: 'right',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'justify',
        }
        return alignment_map.get(alignment, 'left')

    def get_run_style(self, run):
        """Extract styling from a run and return CSS style string"""
        styles = []
        
        if run.bold:
            styles.append("font-weight: bold")
        if run.italic:
            styles.append("font-style: italic")
        if run.underline:
            styles.append("text-decoration: underline")
        
        # Font color
        if run.font.color and run.font.color.rgb:
            try:
                rgb = run.font.color.rgb
                # Handle both RGBColor object and tuple formats
                if hasattr(rgb, 'red'):
                    r, g, b = rgb.red, rgb.green, rgb.blue
                else:
                    r, g, b = rgb[0], rgb[1], rgb[2]
                styles.append(f"color: rgb({r}, {g}, {b})")
            except (AttributeError, IndexError, TypeError):
                # Skip color if we can't extract RGB values
                pass
        
        # Font size
        if run.font.size:
            try:
                # Convert from Pt to pixels (approximate conversion)
                size_px = int(run.font.size.pt * 1.33)
                styles.append(f"font-size: {size_px}px")
            except (AttributeError, TypeError):
                # Skip font size if we can't extract it
                pass
        
        # Font name
        if run.font.name:
            styles.append(f"font-family: '{run.font.name}', sans-serif")
        
        return "; ".join(styles) if styles else ""

    def convert_paragraph_to_html(self, paragraph):
        """Convert a docx paragraph to HTML"""
        if not paragraph.text.strip():
            return "<br/>"
        
        # Get paragraph alignment
        alignment = self.convert_paragraph_alignment(paragraph.alignment)
        
        # Start building the HTML
        html_content = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Escape HTML characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Get run styling
            style = self.get_run_style(run)
            
            if style:
                html_content.append(f'<span style="{style}">{text}</span>')
            else:
                html_content.append(text)
        
        # Wrap in paragraph tag with alignment
        content = ''.join(html_content)
        if alignment != 'left':
            return f'<p style="text-align: {alignment}">{content}</p>'
        else:
            return f'<p>{content}</p>'

    def convert_table_to_html(self, table):
        """Convert a docx table to HTML"""
        html = ['<table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse;">']
        
        for row in table.rows:
            html.append('<tr>')
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_content.append(self.convert_paragraph_to_html(paragraph))
                
                cell_html = ''.join(cell_content) if cell_content else '&nbsp;'
                html.append(f'<td>{cell_html}</td>')
            html.append('</tr>')
        
        html.append('</table>')
        return ''.join(html)

    def convert_docx_to_html(self, docx_path, html_path):
        """Convert a single docx file to HTML"""
        try:
            doc = Document(docx_path)
            
            # Start building HTML
            html_parts = [
                '<!DOCTYPE html>',
                '<html>',
                '<head>',
                '<meta charset="UTF-8">',
                '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
                f'<title>{Path(docx_path).stem}</title>',
                '<style>',
                'body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }',
                'table { margin: 20px 0; width: 100%; }',
                'p { margin: 10px 0; }',
                '</style>',
                '</head>',
                '<body>',
            ]
            
            # Convert document content
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    for para in doc.paragraphs:
                        if para._element == element:
                            html_parts.append(self.convert_paragraph_to_html(para))
                            break
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    for table in doc.tables:
                        if table._element == element:
                            html_parts.append(self.convert_table_to_html(table))
                            break
            
            html_parts.extend(['</body>', '</html>'])
            
            # Write HTML file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_parts))
            
            return True
            
        except Exception as e:
            self.errors.append(f"Error converting {docx_path}: {str(e)}")
            return False

    def find_word_documents(self, directory):
        """Find all .docx files in directory and subdirectories"""
        word_files = []
        directory = Path(directory)
        
        for file_path in directory.rglob('*.docx'):
            # Skip temporary files that start with ~$
            if not file_path.name.startswith('~$'):
                word_files.append(file_path)
        
        return word_files

    def convert_directory(self, directory):
        """Convert all Word documents in directory and subdirectories"""
        directory = Path(directory)
        
        if not directory.exists():
            click.echo(f"Error: Directory '{directory}' does not exist.", err=True)
            return False
        
        if not directory.is_dir():
            click.echo(f"Error: '{directory}' is not a directory.", err=True)
            return False
        
        # Find all Word documents
        word_files = self.find_word_documents(directory)
        
        if not word_files:
            click.echo(f"No Word documents (.docx) found in '{directory}' and its subdirectories.")
            return True
        
        click.echo(f"Found {len(word_files)} Word document(s) to convert...")
        
        # Convert each file
        for docx_path in word_files:
            html_path = docx_path.with_suffix('.html')
            
            click.echo(f"Converting: {docx_path.relative_to(directory)}")
            
            if self.convert_docx_to_html(docx_path, html_path):
                self.converted_count += 1
                click.echo(f"  ✓ Converted to: {html_path.relative_to(directory)}")
            else:
                self.error_count += 1
                click.echo(f"  ✗ Failed to convert", err=True)
        
        return True


@click.command()
@click.argument('directory', type=click.Path(exists=True, file_okay=False, dir_okay=True, readable=True))
@click.option('--verbose', '-v', is_flag=True, help='Show detailed output')
def main(directory, verbose):
    """
    Convert Word documents (.docx) to HTML format.
    
    DIRECTORY: Path to the directory containing Word documents to convert.
    The tool will search recursively through all subdirectories.
    """
    click.echo("🔄 GIACONVERT - Word to HTML Converter")
    click.echo("=" * 40)
    
    converter = WordToHTMLConverter()
    
    # Convert documents
    success = converter.convert_directory(directory)
    
    if not success:
        sys.exit(1)
    
    # Show summary
    click.echo("\n" + "=" * 40)
    click.echo("📊 Conversion Summary:")
    click.echo(f"  ✅ Successfully converted: {converter.converted_count}")
    click.echo(f"  ❌ Failed conversions: {converter.error_count}")
    
    if converter.errors and (verbose or converter.error_count > 0):
        click.echo("\n🚨 Errors encountered:")
        for error in converter.errors:
            click.echo(f"  • {error}")
    
    if converter.converted_count > 0:
        click.echo(f"\n🎉 Conversion completed! HTML files saved in the same directories as the original Word documents.")


if __name__ == '__main__':
    main()