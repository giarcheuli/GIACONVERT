#!/usr/bin/env python3
"""
Debug script to explore headers and footers in Word documents
"""

from docx import Document
import sys

def explore_headers_footers(docx_path):
    """Explore headers and footers in a Word document"""
    print(f"Exploring headers and footers in: {docx_path}")
    
    doc = Document(docx_path)
    
    print(f"Number of sections: {len(doc.sections)}")
    
    for i, section in enumerate(doc.sections):
        print(f"\n=== Section {i + 1} ===")
        
        # Check if section has header
        if section.header:
            print("Header found:")
            for j, paragraph in enumerate(section.header.paragraphs):
                print(f"  Header paragraph {j}: '{paragraph.text}'")
                if paragraph.text.strip():
                    for k, run in enumerate(paragraph.runs):
                        print(f"    Run {k}: '{run.text}' (bold: {run.bold}, italic: {run.italic})")
        else:
            print("No header found")
            
        # Check if section has footer
        if section.footer:
            print("Footer found:")
            for j, paragraph in enumerate(section.footer.paragraphs):
                print(f"  Footer paragraph {j}: '{paragraph.text}'")
                if paragraph.text.strip():
                    for k, run in enumerate(paragraph.runs):
                        print(f"    Run {k}: '{run.text}' (bold: {run.bold}, italic: {run.italic})")
        else:
            print("No footer found")
            
        # Check for first page header/footer
        try:
            if hasattr(section, 'first_page_header') and section.first_page_header:
                print("First page header found:")
                for j, paragraph in enumerate(section.first_page_header.paragraphs):
                    print(f"  First header paragraph {j}: '{paragraph.text}'")
        except AttributeError:
            print("First page header not accessible")
            
        try:
            if hasattr(section, 'first_page_footer') and section.first_page_footer:
                print("First page footer found:")
                for j, paragraph in enumerate(section.first_page_footer.paragraphs):
                    print(f"  First footer paragraph {j}: '{paragraph.text}'")
        except AttributeError:
            print("First page footer not accessible")
            
        # Check for even page header/footer
        try:
            if hasattr(section, 'even_page_header') and section.even_page_header:
                print("Even page header found:")
                for j, paragraph in enumerate(section.even_page_header.paragraphs):
                    print(f"  Even header paragraph {j}: '{paragraph.text}'")
        except AttributeError:
            print("Even page header not accessible")
            
        try:
            if hasattr(section, 'even_page_footer') and section.even_page_footer:
                print("Even page footer found:")
                for j, paragraph in enumerate(section.even_page_footer.paragraphs):
                    print(f"  Even footer paragraph {j}: '{paragraph.text}'")
        except AttributeError:
            print("Even page footer not accessible")

def main():
    if len(sys.argv) > 1:
        explore_headers_footers(sys.argv[1])
    else:
        # Try to explore existing test documents
        test_files = [
            'test_documents/sample_document.docx',
            'test_documents/sample_document_with_images.docx'
        ]
        
        for test_file in test_files:
            try:
                explore_headers_footers(test_file)
                print("\n" + "="*50 + "\n")
            except FileNotFoundError:
                print(f"Test file not found: {test_file}")

if __name__ == '__main__':
    main()