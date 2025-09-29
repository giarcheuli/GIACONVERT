#!/usr/bin/env python3
"""
Test script to create a sample Word document for testing the converter
"""

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_test_document():
    """Create a test Word document with various formatting"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Test Document', 0)
    
    # Introduction paragraph
    intro = doc.add_paragraph('This is a ')
    intro.add_run('test document').bold = True
    intro.add_run(' created to verify the ')
    italic_run = intro.add_run('Word to HTML converter')
    italic_run.italic = True
    italic_run.font.color.rgb = RGBColor(0, 100, 200)
    intro.add_run(' functionality.')
    
    # Formatted paragraph with different alignments
    center_para = doc.add_paragraph('This paragraph is centered')
    center_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # List-like content
    doc.add_paragraph('Features being tested:')
    doc.add_paragraph('• Bold and italic text formatting', style='List Bullet')
    doc.add_paragraph('• Color changes in text', style='List Bullet')
    doc.add_paragraph('• Different paragraph alignments', style='List Bullet')
    doc.add_paragraph('• Table creation and formatting', style='List Bullet')
    
    # Create a simple table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Fill table with data
    headers = ['Name', 'Age', 'City']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # Sample data
    data = [
        ['John Doe', '30', 'New York'],
        ['Jane Smith', '25', 'Los Angeles']
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.cell(row_idx, col_idx).text = cell_data
    
    # Final paragraph
    conclusion = doc.add_paragraph()
    conclusion.add_run('This document demonstrates various ').font.size = 14
    underlined_run = conclusion.add_run('formatting options')
    underlined_run.underline = True
    underlined_run.font.color.rgb = RGBColor(255, 0, 0)
    conclusion.add_run(' that should be preserved in the HTML conversion.')
    
    return doc

def main():
    # Create test directory
    test_dir = 'test_documents'
    os.makedirs(test_dir, exist_ok=True)
    
    # Create and save test document
    doc = create_test_document()
    doc_path = os.path.join(test_dir, 'sample_document.docx')
    doc.save(doc_path)
    
    print(f"✅ Test document created: {doc_path}")
    print("You can now test GIACONVERT with:")
    print(f"./giaconvert {test_dir}")

if __name__ == '__main__':
    main()